from docx import Document
from pathlib import Path
from typing import Dict, List


class WordParser:
    
    def parse(self, file_path: str) -> Dict:
        doc = Document(file_path)
        
        paragraphs = []
        full_text = []
        
        for para in doc.paragraphs:
            if para.text.strip():
                paragraphs.append({
                    "text": para.text,
                    "style": para.style.name
                })
                full_text.append(para.text)
        
        tables_data = []
        for table in doc.tables:
            table_text = []
            for row in table.rows:
                row_text = [cell.text for cell in row.cells]
                table_text.append(row_text)
            tables_data.append(table_text)
        
        return {
            "file_name": Path(file_path).name,
            "file_type": "docx",
            "total_paragraphs": len(paragraphs),
            "total_tables": len(tables_data),
            "full_text": "\n\n".join(full_text),
            "paragraphs": paragraphs,
            "tables": tables_data,
            "metadata": {
                "title": doc.core_properties.title or "",
                "author": doc.core_properties.author or "",
                "subject": doc.core_properties.subject or "",
                "created": str(doc.core_properties.created) if doc.core_properties.created else ""
            }
        }
    
    def extract_tables(self, file_path: str) -> List[List[List[str]]]:
        doc = Document(file_path)
        tables = []
        
        for table in doc.tables:
            table_data = []
            for row in table.rows:
                row_data = [cell.text for cell in row.cells]
                table_data.append(row_data)
            tables.append(table_data)
        
        return tables
